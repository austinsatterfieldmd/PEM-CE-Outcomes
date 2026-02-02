"""
Outcomes Document Parser - Extract questions from Word templates.

Parses the PER Med Affairs Outcomes Questions Review template format to extract:
- Question stems and answer options
- Correct answer (highlighted in yellow)
- Learning objectives
- Educational gaps
- Faculty presenter
- References
- CME level (Knowledge vs Competence)
"""
import re
import logging
from dataclasses import dataclass, field
from typing import Optional
from pathlib import Path

try:
    from docx import Document
    from docx.shared import RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

# Word XML namespace
WORD_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def _extract_cell_text_with_list_markers(cell) -> str:
    """
    Extract text from a cell, including list markers (A., B., etc.) that
    Word stores in formatting rather than text.
    """
    result_parts = []
    list_item_index = 0

    for para in cell.paragraphs:
        text = para.text.strip()

        # Check if this paragraph is part of a numbered/lettered list
        numPr = para._element.find(f'.//{WORD_NS}numPr')
        if numPr is not None and text:
            # This is a list item - prepend letter marker
            letter = chr(ord('A') + list_item_index)
            result_parts.append(f"{letter}. {text}")
            list_item_index += 1
        elif text:
            result_parts.append(text)

    return '\n'.join(result_parts)


@dataclass
class ParsedQuestion:
    """A question extracted from the outcomes document."""
    question_number: int
    question_stem: str
    options: list[str]
    correct_answer: str  # "A", "B", "C", "D", or "E"
    correct_answer_text: str
    learning_objective: str
    educational_gap: str
    faculty_presenter: str
    content_section: str
    references: list[str]
    cme_level: str  # "Knowledge" or "Competence"
    raw_table_data: dict = field(default_factory=dict)

    def to_dict(self) -> dict:
        """Convert to dictionary for JSON serialization."""
        return {
            "question_number": self.question_number,
            "question_stem": self.question_stem,
            "options": self.options,
            "correct_answer": self.correct_answer,
            "correct_answer_text": self.correct_answer_text,
            "learning_objective": self.learning_objective,
            "educational_gap": self.educational_gap,
            "faculty_presenter": self.faculty_presenter,
            "content_section": self.content_section,
            "references": self.references,
            "cme_level": self.cme_level,
        }


@dataclass
class ParsedDocument:
    """Complete parsed outcomes document."""
    filename: str
    activity_title: str
    questions: list[ParsedQuestion]
    parse_warnings: list[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        """Convert to dictionary for JSON serialization."""
        return {
            "filename": self.filename,
            "activity_title": self.activity_title,
            "question_count": len(self.questions),
            "questions": [q.to_dict() for q in self.questions],
            "parse_warnings": self.parse_warnings,
        }


def _extract_question_number(text: str) -> Optional[int]:
    """Extract question number from text like '1. The phase 3...'"""
    match = re.match(r'^(\d+)\.\s*', text.strip())
    if match:
        return int(match.group(1))
    return None


def _extract_options_from_text(text: str) -> tuple[str, list[str]]:
    """
    Extract question stem and options from combined text.

    Handles formats like:
    - "1. Question stem? A. Option A B. Option B..."
    - "1. Question stem?\nA) Option A\nB) Option B..."
    - "1. Question stem?\n\nOption A\nOption B..." (no letter prefixes)
    """
    # Normalize option markers (A. or A) or A:)
    text = re.sub(r'\n?([A-E])\)\s*', r'\n\1. ', text)
    text = re.sub(r'\n?([A-E]):\s*', r'\n\1. ', text)

    # Split on option markers
    pattern = r'\n([A-E])\.\s+'
    parts = re.split(pattern, text)

    if len(parts) < 3:
        # Try without newlines
        pattern = r'([A-E])\.\s+'
        parts = re.split(pattern, text)

    if len(parts) >= 3:
        # First part is the stem
        stem = parts[0].strip()
        # Remove question number from stem
        stem = re.sub(r'^\d+\.\s*', '', stem)

        # Reconstruct options (clean "(best answer)" marker for final output)
        options = []
        for i in range(1, len(parts), 2):
            if i + 1 < len(parts):
                letter = parts[i]
                option_text = parts[i + 1].strip()
                # Keep original for correct answer detection, clean for output
                options.append(f"{letter}. {option_text}")

        return stem, options

    # Fallback: Try splitting on double newlines (options without letter prefixes)
    # This handles formats like "stem?\n\nOption 1\nOption 2\nOption 3\nOption 4"
    if '\n\n' in text:
        parts = text.split('\n\n', 1)
        if len(parts) == 2:
            stem = parts[0].strip()
            stem = re.sub(r'^\d+\.\s*', '', stem)

            # Split remaining text into options by single newlines
            option_lines = [line.strip() for line in parts[1].split('\n') if line.strip()]

            if len(option_lines) >= 2:
                # Assign letters A, B, C, D, E automatically
                letters = 'ABCDE'
                options = []
                for i, opt_text in enumerate(option_lines[:5]):  # Max 5 options
                    options.append(f"{letters[i]}. {opt_text}")
                return stem, options

    # Final fallback: return as-is with no options
    stem = re.sub(r'^\d+\.\s*', '', text.strip())
    return stem, []


def _find_correct_answer_by_highlight(table, question_row_idx: int) -> Optional[str]:
    """
    Find the correct answer by looking for yellow highlighting.
    Returns the letter (A, B, C, D, E) of the highlighted option.
    """
    if not DOCX_AVAILABLE:
        return None

    try:
        # Look through the question/answer cell for highlighting
        for row in table.rows[question_row_idx:question_row_idx + 2]:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # Check for yellow highlight
                        if run.font.highlight_color:
                            text = run.text.strip()
                            # Check if this starts with an option letter
                            match = re.match(r'^([A-E])[\.\)\:]?\s*', text)
                            if match:
                                return match.group(1)
    except Exception as e:
        logger.debug(f"Error checking highlights: {e}")

    return None


def _determine_cme_level(table_data: dict) -> str:
    """Determine if question is Knowledge or Competence based on table data."""
    level_text = table_data.get("cme_level", "").lower()

    if "competence" in level_text:
        return "Competence"
    elif "knowledge" in level_text:
        return "Knowledge"

    # Check for checkbox-style format "Knowledge\n Competence" where checked has different format
    # Default to Knowledge if unclear
    return "Knowledge"


def _parse_table_to_dict(table) -> dict:
    """Parse a question analysis table into a dictionary."""
    data = {}

    for row in table.rows:
        if len(row.cells) < 2:
            continue

        key = row.cells[0].text.strip().lower()

        # Map common field names
        if "faculty" in key or "presenter" in key:
            data["faculty"] = row.cells[1].text.strip()
        elif "content section" in key or "slide" in key or "module" in key:
            data["content_section"] = row.cells[1].text.strip()
        elif "educational gap" in key:
            data["educational_gap"] = row.cells[1].text.strip()
        elif "learning objective" in key:
            data["learning_objective"] = row.cells[1].text.strip()
        elif "outcomes question" in key:
            # Use special extraction to capture list markers (A., B., etc.)
            data["question_and_answers"] = _extract_cell_text_with_list_markers(row.cells[1])
        elif "tagging" in key:
            data["tagging"] = row.cells[1].text.strip()
        elif "reference" in key:
            data["references"] = row.cells[1].text.strip()
        elif "knowledge" in key.lower() and "competence" in key.lower():
            data["cme_level"] = row.cells[1].text.strip()
        elif "pre/post" in key or "data" in key:
            data["pre_post_data"] = row.cells[1].text.strip()
        elif "summary" in key or "recommendation" in key:
            data["summary"] = row.cells[1].text.strip()

    return data


def _split_table_into_question_sections(table) -> list[list]:
    """
    Split a table into multiple question sections.

    Some documents have multiple questions within a single table,
    separated by "INDIVIDUAL QUESTION ANALYSIS" header rows.

    Returns:
        List of row groups, where each group is one question's rows
    """
    sections = []
    current_section = []

    for row in table.rows:
        first_cell = row.cells[0].text.strip().lower() if row.cells else ""

        # Check if this is a header row (start of new question)
        if "individual question analysis" in first_cell:
            # Save previous section if it has content
            if current_section:
                sections.append(current_section)
            current_section = []
        else:
            current_section.append(row)

    # Don't forget the last section
    if current_section:
        sections.append(current_section)

    return sections


def _parse_rows_to_dict(rows) -> dict:
    """Parse a list of table rows into a dictionary."""
    data = {}

    for row in rows:
        if len(row.cells) < 2:
            continue

        key = row.cells[0].text.strip().lower()

        # Map common field names
        if "faculty" in key or "presenter" in key:
            data["faculty"] = row.cells[1].text.strip()
        elif "content section" in key or "slide" in key or "module" in key:
            data["content_section"] = row.cells[1].text.strip()
        elif "educational gap" in key:
            data["educational_gap"] = row.cells[1].text.strip()
        elif "learning objective" in key:
            data["learning_objective"] = row.cells[1].text.strip()
        elif "outcomes question" in key:
            # Use special extraction to capture list markers (A., B., etc.)
            data["question_and_answers"] = _extract_cell_text_with_list_markers(row.cells[1])
        elif "tagging" in key:
            data["tagging"] = row.cells[1].text.strip()
        elif "reference" in key:
            data["references"] = row.cells[1].text.strip()
        elif "knowledge" in key.lower() and "competence" in key.lower():
            data["cme_level"] = row.cells[1].text.strip()
        elif "pre/post" in key or "data" in key:
            data["pre_post_data"] = row.cells[1].text.strip()
        elif "summary" in key or "recommendation" in key:
            data["summary"] = row.cells[1].text.strip()

    return data


def _infer_correct_answer(question_text: str, options: list[str]) -> Optional[str]:
    """
    Attempt to infer correct answer from context clues.
    Returns None if cannot determine.
    """
    # First check for "(best answer)" marker in options
    for opt in options:
        if "(best answer)" in opt.lower():
            # Extract the letter from the option (e.g., "A. Some text (best answer)")
            match = re.match(r'^([A-E])[\.\):]?\s*', opt)
            if match:
                return match.group(1).upper()

    # Look for patterns like "correct answer is A" or "answer: B"
    patterns = [
        r'correct\s+answer\s*(?:is|:)?\s*([A-E])',
        r'answer\s*:\s*([A-E])',
        r'\*\s*([A-E])\.',  # Asterisk before option
    ]

    full_text = question_text + " " + " ".join(options)
    for pattern in patterns:
        match = re.search(pattern, full_text, re.IGNORECASE)
        if match:
            return match.group(1).upper()

    return None


def _clean_option_text(option: str) -> str:
    """Remove (best answer) marker from option text."""
    # Remove variations of "(best answer)" marker
    cleaned = re.sub(r'\s*\(best\s+answer\)\s*', '', option, flags=re.IGNORECASE)
    return cleaned.strip()


def parse_outcomes_document(file_path: str) -> ParsedDocument:
    """
    Parse an outcomes document and extract all questions.

    Args:
        file_path: Path to the Word document

    Returns:
        ParsedDocument with all extracted questions
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    doc = Document(file_path)
    questions = []
    warnings = []
    activity_title = ""

    # Extract activity title from first paragraph or header
    for para in doc.paragraphs[:5]:
        text = para.text.strip()
        if text and "outcome" not in text.lower() and "note" not in text.lower():
            activity_title = text
            break

    # Process each table
    question_number = 0
    for table_idx, table in enumerate(doc.tables):
        # Skip non-question tables (header table, reviewer table)
        first_cell_text = table.rows[0].cells[0].text.strip().lower() if table.rows else ""

        if "individual question analysis" not in first_cell_text:
            continue

        # Split table into question sections (handles multi-question tables)
        question_sections = _split_table_into_question_sections(table)

        if not question_sections:
            # Fallback: treat entire table as one question
            question_sections = [list(table.rows)]

        for section_idx, section_rows in enumerate(question_sections):
            # Parse this question section
            table_data = _parse_rows_to_dict(section_rows)

            question_text = table_data.get("question_and_answers", "")
            if not question_text:
                warnings.append(f"Table {table_idx + 1}, Section {section_idx + 1}: No question text found")
                continue

            question_number += 1

            # Extract stem and options
            stem, options = _extract_options_from_text(question_text)

            if not options:
                warnings.append(f"Question {question_number}: Could not parse answer options")

            # Find correct answer - try to infer from text patterns (uses raw options with markers)
            correct_answer = _infer_correct_answer(question_text, options)
            if not correct_answer:
                correct_answer = "A"  # Default
                warnings.append(f"Question {question_number}: Could not determine correct answer, defaulting to A")

            # Clean "(best answer)" from options for storage
            cleaned_options = [_clean_option_text(opt) for opt in options]

            # Get correct answer text (from cleaned options)
            correct_answer_text = ""
            for opt in cleaned_options:
                if opt.startswith(f"{correct_answer}.") or opt.startswith(f"{correct_answer})"):
                    correct_answer_text = opt[2:].strip() if len(opt) > 2 else opt
                    break

            # Parse references
            references_text = table_data.get("references", "")
            references = [r.strip() for r in re.split(r'[;\n]', references_text) if r.strip()]

            # Create parsed question
            parsed_q = ParsedQuestion(
                question_number=question_number,
                question_stem=stem,
                options=cleaned_options,
                correct_answer=correct_answer,
                correct_answer_text=correct_answer_text,
                learning_objective=table_data.get("learning_objective", ""),
                educational_gap=table_data.get("educational_gap", ""),
                faculty_presenter=table_data.get("faculty", ""),
                content_section=table_data.get("content_section", ""),
                references=references,
                cme_level=_determine_cme_level(table_data),
                raw_table_data=table_data,
            )
            questions.append(parsed_q)

    return ParsedDocument(
        filename=path.name,
        activity_title=activity_title,
        questions=questions,
        parse_warnings=warnings,
    )


def parse_outcomes_document_from_bytes(file_bytes: bytes, filename: str) -> ParsedDocument:
    """
    Parse an outcomes document from bytes (for uploaded files).

    Args:
        file_bytes: The document content as bytes
        filename: Original filename

    Returns:
        ParsedDocument with all extracted questions
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    import io
    doc = Document(io.BytesIO(file_bytes))

    questions = []
    warnings = []
    activity_title = ""

    # Extract activity title from first paragraph or header
    for para in doc.paragraphs[:5]:
        text = para.text.strip()
        if text and "outcome" not in text.lower() and "note" not in text.lower():
            activity_title = text
            break

    # Process each table
    question_number = 0
    for table_idx, table in enumerate(doc.tables):
        # Skip non-question tables
        first_cell_text = table.rows[0].cells[0].text.strip().lower() if table.rows else ""

        if "individual question analysis" not in first_cell_text:
            continue

        # Split table into question sections (handles multi-question tables)
        question_sections = _split_table_into_question_sections(table)

        if not question_sections:
            # Fallback: treat entire table as one question
            question_sections = [list(table.rows)]

        for section_idx, section_rows in enumerate(question_sections):
            # Parse this question section
            table_data = _parse_rows_to_dict(section_rows)

            question_text = table_data.get("question_and_answers", "")
            if not question_text:
                warnings.append(f"Table {table_idx + 1}, Section {section_idx + 1}: No question text found")
                continue

            question_number += 1

            # Extract stem and options
            stem, options = _extract_options_from_text(question_text)

            if not options:
                warnings.append(f"Question {question_number}: Could not parse answer options")

            # Find correct answer - try to infer from text patterns (uses raw options with markers)
            correct_answer = _infer_correct_answer(question_text, options)
            if not correct_answer:
                correct_answer = "A"
                warnings.append(f"Question {question_number}: Could not determine correct answer, defaulting to A")

            # Clean "(best answer)" from options for storage
            cleaned_options = [_clean_option_text(opt) for opt in options]

            # Get correct answer text (from cleaned options)
            correct_answer_text = ""
            for opt in cleaned_options:
                if opt.startswith(f"{correct_answer}.") or opt.startswith(f"{correct_answer})"):
                    correct_answer_text = opt[2:].strip() if len(opt) > 2 else opt
                    break

            # Parse references
            references_text = table_data.get("references", "")
            references = [r.strip() for r in re.split(r'[;\n]', references_text) if r.strip()]

            parsed_q = ParsedQuestion(
                question_number=question_number,
                question_stem=stem,
                options=cleaned_options,
                correct_answer=correct_answer,
                correct_answer_text=correct_answer_text,
                learning_objective=table_data.get("learning_objective", ""),
                educational_gap=table_data.get("educational_gap", ""),
                faculty_presenter=table_data.get("faculty", ""),
                content_section=table_data.get("content_section", ""),
                references=references,
                cme_level=_determine_cme_level(table_data),
                raw_table_data=table_data,
            )
            questions.append(parsed_q)

    return ParsedDocument(
        filename=filename,
        activity_title=activity_title,
        questions=questions,
        parse_warnings=warnings,
    )
